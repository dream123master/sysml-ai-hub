#!/usr/bin/env python3
"""Generate the downloadable SysML AI research report from data/catalog.json."""

from __future__ import annotations

import json
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CATALOG_PATH = ROOT / "data" / "catalog.json"
OUTPUT_PATH = ROOT / "public" / "reports" / "sysml-ai-research-report.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "007079"
MUTED = "5D6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BODY_FONT = "Microsoft YaHei"
BODY_SIZE = 11


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_column_widths(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("SysML AI Hub · 调研报告")
    set_run_font(left, size=9, color=MUTED, bold=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("第 ")
    set_run_font(prefix, size=9, color=MUTED)
    add_field(paragraph, "PAGE")
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def add_title_block(document: Document, data: dict) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(6)
    run = kicker.add_run("SYSML AI ECOSYSTEM RESEARCH")
    set_run_font(run, size=9, color=TEAL, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("SysML 大模型插件与 Skill 调研报告")
    set_run_font(run, size=25, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("覆盖 SysML v1 / SysML v2、MCP Server、Agent、Skill 与相关开源工具")
    set_run_font(run, size=11.5, color=MUTED)

    meta = document.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    values = [
        ("报告更新时间", data["meta"]["lastUpdated"]),
        ("下次计划更新", data["meta"]["nextUpdate"]),
        ("收录项目", f'{len(data["projects"])} 个'),
        ("自动更新周期", f'每 {data["meta"]["intervalDays"]} 天'),
    ]
    for index, (label, value) in enumerate(values):
        cell = meta.cell(index // 2, index % 2)
        cell.text = ""
        label_run = cell.paragraphs[0].add_run(f"{label}\n")
        set_run_font(label_run, size=8.5, color=MUTED, bold=True)
        value_run = cell.paragraphs[0].add_run(value)
        set_run_font(value_run, size=11, color=NAVY, bold=True)
        set_cell_shading(cell, LIGHT_GRAY if index > 1 else LIGHT_BLUE)
    set_column_widths(meta, [4680, 4680])


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(text), size=11, color=NAVY)


def add_overview(document: Document, data: dict) -> None:
    document.add_heading("1. 调研概览", level=1)
    verified = sum(p["verification"] == "verified" for p in data["projects"])
    open_source = sum(p["openSource"] == "open" for p in data["projects"])
    version_counts = Counter(v for p in data["projects"] for v in p["sysmlVersions"])
    paragraph = document.add_paragraph()
    paragraph.add_run("本报告与网页目录使用同一份数据源。")
    set_run_font(paragraph.runs[0], color=NAVY, bold=True)
    run = paragraph.add_run(
        f" 当前共收录 {len(data['projects'])} 个项目，其中 {verified} 个已验证、{open_source} 个确认为开源。"
        " 调研重点是这些工具在系统建模生命周期中的实际用途，而不只是在何种 AI 客户端中运行。"
    )
    set_run_font(run, color=NAVY)
    add_bullet(document, "标准覆盖：" + "；".join(f"{name}（{count}）" for name, count in version_counts.items()))
    add_bullet(document, "数据来源：" + "、".join(data["meta"]["sources"]))
    add_bullet(document, "收录方法：" + data["meta"]["methodology"]["zh"])


def add_comparison(document: Document, data: dict) -> None:
    document.add_heading("2. 同类型工具对比", level=1)
    document.add_paragraph("下表按工具形态汇总，便于快速判断不同类型在 SysML 工作流中的定位。")
    groups = defaultdict(list)
    for project in data["projects"]:
        groups[project["kind"]].append(project)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["类型", "代表项目", "主要适用阶段", "典型能力"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
        set_cell_shading(cell, LIGHT_BLUE)
        set_run_font(cell.paragraphs[0].runs[0], size=9, color=NAVY, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for kind, projects in sorted(groups.items()):
        cells = table.add_row().cells
        use_cases = list(dict.fromkeys(x for p in projects for x in p["useCases"]["zh"]))[:4]
        capabilities = list(dict.fromkeys(x for p in projects for x in p["capabilities"]["zh"]))[:4]
        values = [kind, "、".join(p["name"] for p in projects), "、".join(use_cases), "、".join(capabilities)]
        for cell, value in zip(cells, values):
            cell.text = value
            set_run_font(cell.paragraphs[0].runs[0], size=8.5, color=NAVY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_column_widths(table, [1260, 2700, 2700, 2700])


def add_project_details(document: Document, data: dict) -> None:
    document.add_heading("3. 项目明细", level=1)
    for index, project in enumerate(data["projects"], 1):
        heading = document.add_heading(f"3.{index} {project['name']}", level=2)
        heading.paragraph_format.page_break_before = index > 1
        tags = " · ".join([
            project["kind"],
            " / ".join(project["sysmlVersions"]),
            "已验证" if project["verification"] == "verified" else "待验证",
            "开源" if project["openSource"] == "open" else "闭源" if project["openSource"] == "closed" else "开源状态未确认",
        ])
        tag_paragraph = document.add_paragraph()
        tag_paragraph.paragraph_format.space_after = Pt(8)
        set_run_font(tag_paragraph.add_run(tags), size=9, color=TEAL, bold=True)

        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        rows = [
            ("功能说明", project["functionality"]["zh"]),
            ("适用场景", "、".join(project["useCases"]["zh"])),
            ("核心能力", "、".join(project["capabilities"]["zh"])),
            ("支持平台", "、".join(project["platforms"])),
            ("隐私与离线", f"{project['privacy']['zh']}（离线能力：{project['offline']}）"),
            ("项目地址", project.get("githubUrl") or project["homepageUrl"]),
            ("核验来源", project["sourceUrl"]),
        ]
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
            set_cell_shading(cells[0], LIGHT_BLUE)
            set_run_font(cells[0].paragraphs[0].runs[0], size=9, color=NAVY, bold=True)
            set_run_font(cells[1].paragraphs[0].runs[0], size=9, color=NAVY)
            cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_column_widths(table, [1700, 7660])


def add_notes(document: Document, data: dict) -> None:
    document.add_heading("4. 使用说明与局限", level=1)
    add_bullet(document, "“待验证”项目由公开信息自动发现，功能、隐私、安装复杂度等字段仍需人工复核。")
    add_bullet(document, "GitHub 星标、最近活动时间及项目状态可能随下一次自动检查而变化。")
    add_bullet(document, "本报告用于技术调研与选型初筛，不替代对许可证、安全策略和实际建模效果的独立验证。")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    set_run_font(paragraph.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}（Asia/Shanghai）"), size=9, color=MUTED)


def build_report() -> None:
    data = json.loads(CATALOG_PATH.read_text(encoding="utf-8"))
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_page(document)
    configure_styles(document)
    add_title_block(document, data)
    add_overview(document, data)
    add_comparison(document, data)
    add_project_details(document, data)
    add_notes(document, data)
    document.core_properties.title = "SysML 大模型插件与 Skill 调研报告"
    document.core_properties.subject = "SysML v1/v2 AI 插件、MCP Server、Agent 与 Skill 调研"
    document.core_properties.author = "SysML AI Hub"
    document.save(OUTPUT_PATH)
    print(f"Generated {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
